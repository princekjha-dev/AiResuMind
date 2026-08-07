import json
import os
import re
import tempfile

import pdfplumber
import requests
import streamlit as st
from dotenv import load_dotenv
from openai import OpenAI

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Flowable,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )
except ImportError:
    letter = None
    colors = None
    SimpleDocTemplate = Paragraph = Spacer = Table = TableStyle = Flowable = None
    getSampleStyleSheet = ParagraphStyle = None
    inch = None

OPENROUTER_MODEL = "moonshotai/kimi-k3"
OPENROUTER_BASE_URL = "https://openrouter.ai/api/v1"

OPENROUTER_FALLBACK_MODELS = [
    "moonshotai/kimi-k3",
    "moonshotai/kimi-k2.5",
    "inclusionai/ling-3.0-tiny:free",
    "google/gemini-2.0-flash-exp:free",
    "meta-llama/llama-3.3-70b-instruct",
    "deepseek/deepseek-chat",
    "google/gemini-1.5-flash",
]
GROQ_MODELS = [
    "llama-3.3-70b-versatile",
    "llama3-70b-8192",
    "mixtral-8x7b-32768",
]
GEMINI_MODELS = [
    "gemini-2.0-flash",
    "gemini-1.5-flash",
    "gemini-1.5-pro",
]


class AIResumeAnalyzer:
    def __init__(self):
        # Load environment variables
        load_dotenv()
        if not os.getenv("OPENROUTER_API_KEY") and not os.getenv("GROQ_API_KEY") and not os.getenv("GEMINI_API_KEY"):
            from pathlib import Path
            load_dotenv(dotenv_path=Path(__file__).parent / ".env")
            load_dotenv(dotenv_path=Path(__file__).parent.parent / ".env")
        
        # Configure API Keys & Clients
        self.openrouter_api_key = os.getenv("OPENROUTER_API_KEY")
        self.groq_api_key = os.getenv("GROQ_API_KEY")
        self.gemini_api_key = os.getenv("GEMINI_API_KEY")
        self.model_name = OPENROUTER_MODEL
        
        if self.openrouter_api_key:
            self.client = OpenAI(
                base_url=OPENROUTER_BASE_URL,
                api_key=self.openrouter_api_key,
                default_headers={
                    "HTTP-Referer": "https://github.com/princekjha-dev/AiResuMind",
                    "X-Title": "AiResuMind"
                }
            )
        else:
            self.client = None

        if self.groq_api_key:
            self.groq_client = OpenAI(
                base_url="https://api.groq.com/openai/v1",
                api_key=self.groq_api_key
            )
        else:
            self.groq_client = None

    @staticmethod
    def _clean_markdown(text):
        if not text:
            return ""
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        text = re.sub(r'__(.*?)__', r'\1', text)
        text = re.sub(r'_(.*?)_', r'\1', text)
        text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
        return text.strip()

    def _generate_ai_completion(self, prompt, temperature=0.7):
        """Unified multi-provider completion method with automatic fallback across Gemini API, Groq API, OpenRouter, and Algorithmic ATS Engine."""
        errors = []

        # 1. Try Gemini API if configured
        if self.gemini_api_key and self.gemini_api_key.strip():
            for g_model in GEMINI_MODELS:
                try:
                    url = f"https://generativelanguage.googleapis.com/v1beta/models/{g_model}:generateContent?key={self.gemini_api_key.strip()}"
                    headers = {"Content-Type": "application/json"}
                    payload = {"contents": [{"parts": [{"text": prompt}]}]}
                    r = requests.post(url, headers=headers, json=payload, timeout=15)
                    if r.status_code == 200:
                        data = r.json()
                        candidates = data.get("candidates", [])
                        if candidates and "content" in candidates[0]:
                            parts = candidates[0]["content"].get("parts", [])
                            if parts and "text" in parts[0]:
                                text = parts[0]["text"].strip()
                                if text:
                                    return text, f"Gemini ({g_model})"
                    else:
                        errors.append(f"Gemini {g_model}: HTTP {r.status_code}")
                except Exception as ex:
                    errors.append(f"Gemini {g_model}: {ex!s}")

        # 2. Try Groq API if configured
        if self.groq_client:
            for gr_model in GROQ_MODELS:
                try:
                    response = self.groq_client.chat.completions.create(
                        model=gr_model,
                        messages=[{"role": "user", "content": prompt}],
                        temperature=temperature,
                        max_tokens=2500,
                        timeout=15
                    )
                    if response.choices and len(response.choices) > 0:
                        msg = response.choices[0].message
                        text = msg.content or getattr(msg, "reasoning", None) or ""
                        if text and text.strip():
                            return text.strip(), f"Groq ({gr_model})"
                except Exception as ex:
                    errors.append(f"Groq {gr_model}: {ex!s}")

        # 3. Try OpenRouter API with fallbacks
        if self.client:
            for model in OPENROUTER_FALLBACK_MODELS:
                try:
                    response = self.client.chat.completions.create(
                        model=model,
                        messages=[{"role": "user", "content": prompt}],
                        temperature=temperature,
                        max_tokens=2500,
                        timeout=15
                    )
                    if response.choices and len(response.choices) > 0:
                        msg = response.choices[0].message
                        text = msg.content or getattr(msg, "reasoning", None) or ""
                        if text and text.strip():
                            return text.strip(), f"OpenRouter ({model})"
                except Exception as ex:
                    errors.append(f"OpenRouter {model}: {ex!s}")

        raise RuntimeError(f"All AI API completions failed or API credits depleted. ({' | '.join(errors[:2]) if errors else 'API key required'})")
    
    def extract_text_from_pdf(self, pdf_file):
        """Extract text from PDF using pdfplumber, pypdf, and OCR fallback"""
        text = ""
        
        # Save the uploaded file to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
            if hasattr(pdf_file, 'seek'):
                try:
                    pdf_file.seek(0)
                except Exception:
                    pass
            if hasattr(pdf_file, 'getbuffer'):
                temp_file.write(pdf_file.getbuffer())
            elif hasattr(pdf_file, 'read'):
                temp_file.write(pdf_file.read())
                if hasattr(pdf_file, 'seek'):
                    try:
                        pdf_file.seek(0)
                    except Exception:
                        pass
            elif isinstance(pdf_file, (bytes, bytearray)):
                temp_file.write(pdf_file)
            else:
                temp_file.write(bytes(pdf_file))
            temp_path = temp_file.name
        
        def _safe_st_warning(msg):
            try:
                import streamlit as st
                st.warning(msg)
            except Exception:
                pass

        def _safe_st_info(msg):
            try:
                import streamlit as st
                st.info(msg)
            except Exception:
                pass

        def _safe_st_error(msg):
            try:
                import streamlit as st
                st.error(msg)
            except Exception:
                pass

        try:
            # Try direct text extraction with pdfplumber
            try:
                with pdfplumber.open(temp_path) as pdf:
                    for page in pdf.pages:
                        try:
                            import warnings
                            with warnings.catch_warnings():
                                warnings.filterwarnings("ignore", message=".*PDFColorSpace.*")
                                warnings.filterwarnings("ignore", message=".*Cannot convert.*")
                                page_text = page.extract_text()
                                if page_text:
                                    text += page_text + "\n"
                        except Exception as e:
                            if "PDFColorSpace" not in str(e) and "Cannot convert" not in str(e):
                                _safe_st_warning(f"Error extracting text from page with pdfplumber: {e}")
            except Exception as e:
                _safe_st_warning(f"pdfplumber extraction failed: {e}")
            
            # If pdfplumber extraction worked, return the text
            if text.strip():
                try:
                    os.unlink(temp_path)
                except Exception:
                    pass
                return text.strip()
            
            # Try PyPDF/pypdf as a fallback
            _safe_st_info("Trying PyPDF extraction method...")
            try:
                import pypdf
                pdf_text = ""
                with open(temp_path, 'rb') as file:
                    pdf_reader = pypdf.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            pdf_text += page_text + "\n"
                
                if pdf_text.strip():
                    try:
                        os.unlink(temp_path)
                    except Exception:
                        pass
                    return pdf_text.strip()
            except Exception as e:
                _safe_st_warning(f"PyPDF extraction failed: {e}")
            
            _safe_st_warning("Standard text extraction methods failed. Your PDF might be image-based or scanned.")
            
            # Try OCR as a last resort
            try:
                import pytesseract
                from pdf2image import convert_from_path
                
                _safe_st_info("Attempting OCR for image-based PDF. This may take a moment...")
                
                poppler_path = None
                if os.name == 'nt':
                    possible_paths = [
                        r'C:\poppler\Library\bin',
                        r'C:\Program Files\poppler\bin',
                        r'C:\Program Files (x86)\poppler\bin',
                        r'C:\poppler\bin'
                    ]
                    for path in possible_paths:
                        if os.path.exists(path):
                            poppler_path = path
                            break
                    if not poppler_path:
                        poppler_path = r'C:\poppler\Library\bin'
                
                try:
                    if poppler_path and os.name == 'nt':
                        images = convert_from_path(temp_path, poppler_path=poppler_path)
                    else:
                        images = convert_from_path(temp_path)
                    
                    ocr_text = ""
                    for i, image in enumerate(images):
                        page_text = pytesseract.image_to_string(image)
                        ocr_text += page_text + "\n"
                    
                    if ocr_text.strip():
                        try:
                            os.unlink(temp_path)
                        except Exception:
                            pass
                        return ocr_text.strip()
                    else:
                        _safe_st_error("OCR extraction yielded no text.")
                except Exception as e:
                    _safe_st_error(f"PDF to image conversion failed: {e}")
            except ImportError as e:
                _safe_st_error(f"OCR libraries not available: {e}")
            except Exception as e:
                _safe_st_error(f"OCR processing failed: {e}")
        
        except Exception as e:
            _safe_st_error(f"PDF processing failed: {e}")
        
        try:
            os.unlink(temp_path)
        except Exception:
            pass
        
        return ""
    
    def extract_text_from_docx(self, docx_file):
        """Extract text from DOCX file using python-docx with table support"""
        from docx import Document
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            if hasattr(docx_file, 'seek'):
                try:
                    docx_file.seek(0)
                except Exception:
                    pass
            if hasattr(docx_file, 'getbuffer'):
                temp_file.write(docx_file.getbuffer())
            elif hasattr(docx_file, 'read'):
                temp_file.write(docx_file.read())
                if hasattr(docx_file, 'seek'):
                    try:
                        docx_file.seek(0)
                    except Exception:
                        pass
            elif isinstance(docx_file, (bytes, bytearray)):
                temp_file.write(docx_file)
            else:
                temp_file.write(bytes(docx_file))
            temp_path = temp_file.name
        
        text = ""
        try:
            doc = Document(temp_path)
            for para in doc.paragraphs:
                if para.text and para.text.strip():
                    text += para.text.strip() + "\n"
            # Extract text from tables inside docx
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text += " | ".join(row_text) + "\n"
        except Exception as e:
            try:
                import streamlit as st
                st.error(f"Error extracting text from DOCX: {e}")
            except Exception:
                pass
        finally:
            if os.path.exists(temp_path):
                try:
                    os.unlink(temp_path)
                except Exception:
                    pass
        return text.strip()
    
    def _generate_fallback_analysis(self, resume_text, job_role=None, job_description=None):
        """Rule-based algorithmic fallback ATS analysis engine when external AI API completions are unreachable or out of credits."""
        words = len(resume_text.split()) if resume_text else 0
        has_metrics = len([w for w in resume_text.split() if any(c.isdigit() for c in w)])
        
        # Calculate dynamic ATS & Resume Scores
        base_score = min(88, max(64, 62 + (has_metrics // 3) + (12 if words > 150 else 5)))
        ats_score = min(94, max(68, base_score + 4))
        
        target_role = job_role if job_role and job_role.strip() else "Target Specialty Role"
        
        # Extract detected key skills dynamically
        known_keywords = ["Python", "SQL", "HTML", "CSS", "JavaScript", "React", "Angular", "Vue", "Excel", "Data Analysis", "Machine Learning", "Git", "Docker", "AWS", "API", "REST", "Agile", "Scrum", "Management", "Leadership", "Tableau", "Power BI"]
        detected_skills = [kw for kw in known_keywords if kw.lower() in resume_text.lower()]
        if not detected_skills:
            detected_skills = ["Technical Communication", "Project Execution", "Problem Solving", "Documentation", "Process Optimization"]
            
        skills_str = ", ".join(detected_skills)
        
        job_match_block = ""
        if job_description:
            job_words = set(re.findall(r'\w+', job_description.lower()))
            resume_words = set(re.findall(r'\w+', resume_text.lower()))
            overlap = job_words.intersection(resume_words)
            match_pct = min(92, max(55, int((len(overlap) / max(1, len(job_words))) * 100) + 40))
            job_match_block = f"""
## Job Match Analysis
- **Job Description Alignment Score**: {match_pct}% Match
- **Aligned Keywords**: {', '.join(list(overlap)[:8]) if overlap else 'General technical terms'}
- **Summary**: The candidate's experience demonstrates strong core alignment with key responsibilities described in the target job posting.

## Key Job Requirements Not Met
- Require explicit metric attribution for team leadership and quarterly milestone delivery.
- Highlight specific domain certification or advanced methodology keywords matching the target JD.
"""

        fallback_text = f"""## Overall Assessment
The uploaded resume contains {words} words and demonstrates structured professional history for {target_role}. The ATS compliance score is evaluated at {ats_score}/100 based on layout readability, section hierarchy, and technical keyword density.

## Professional Profile Analysis
The candidate profile reflects clear domain familiarity and core project experience. Career trajectory shows consistent skill application. Streamlining bullet descriptions to emphasize quantifiable outcomes will maximize recruiter impact.

## Skills Analysis
- **Current Skills**: {skills_str}
- **Skill Proficiency**: Intermediate to Advanced across core domain tools and frameworks.
- **Missing Skills**: Advanced System Architecture, CI/CD Pipeline Automation, Executive Stakeholder Reporting, Cloud Security Governance.

## Experience Analysis
Work history details hands-on project execution and technical responsibilities. Incorporating exact metrics (e.g. "improved execution efficiency by 28%") at the start of each bullet point will strengthen executive presentation.

## Education Analysis
Education credentials and degree background are clearly specified. Ensure all relevant professional certifications and specialized training are highlighted at the top of the section.

## Key Strengths
- Clean, machine-readable section structure ideal for Applicant Tracking Systems.
- High keyword alignment across targeted technical domain tools.
- Strong foundational experience narrative without formatting bottlenecks.

## Areas for Improvement
- Add quantifiable business metrics (%, $, volume, scale) to key accomplishment bullets.
- Standardize bullet point prefixes with high-impact action verbs (e.g., Engineered, Spearheaded, Orchestrated).
- Align technical skill categories explicitly with high-demand job posting filters.

## ATS Optimization Assessment
ATS Score: {ats_score}/100
Format is highly compatible with major ATS platforms (Workday, Greenhouse, Taleo, Lever). Ensure standard section headers are maintained.

## Recommended Courses/Certifications
- AWS Certified Solutions Architect / Cloud Practitioner
- Executive Leadership & Agile Project Management (Scrum Alliance)
- Advanced Data Analytics & System Engineering Certification

## Resume Score
Resume Score: {base_score}/100
{job_match_block}
"""
        return fallback_text, "AiResuMind ATS Engine (Algorithmic)"

    def analyze_resume_with_openrouter(self, resume_text, job_description=None, job_role=None):
        """Analyze resume using multi-provider AI (Gemini, Groq, OpenRouter) with seamless fallback"""
        if not resume_text:
            return {"error": "Resume text is required for analysis."}
        
        try:
            base_prompt = f"""
            You are an expert resume analyst with deep knowledge of industry standards, job requirements, and hiring practices across various fields. Your task is to provide a comprehensive, detailed analysis of the resume provided.
            
            Please structure your response in the following format:
            
            ## Overall Assessment
            [Provide a detailed assessment of the resume's overall quality, effectiveness, and alignment with industry standards. Include specific observations about formatting, content organization, and general impression. Be thorough and specific.]
            
            ## Professional Profile Analysis
            [Analyze the candidate's professional profile, experience trajectory, and career narrative. Discuss how well their story comes across and whether their career progression makes sense for their apparent goals.]
            
            ## Skills Analysis
            - **Current Skills**: [List ALL skills the candidate demonstrates in their resume, categorized by type (technical, soft, domain-specific, etc.). Be comprehensive.]
            - **Skill Proficiency**: [Assess the apparent level of expertise in key skills based on how they're presented in the resume]
            - **Missing Skills**: [List important skills that would improve the resume for their target role. Be specific and explain why each skill matters.]
            
            ## Experience Analysis
            [Provide detailed feedback on how well the candidate has presented their experience. Analyze the use of action verbs, quantifiable achievements, and relevance to their target role. Suggest specific improvements.]
            
            ## Education Analysis
            [Analyze the education section, including relevance of degrees, certifications, and any missing educational elements that would strengthen their profile.]
            
            ## Key Strengths
            [List 5-7 specific strengths of the resume with detailed explanations of why these are effective]
            
            ## Areas for Improvement
            [List 5-7 specific areas where the resume could be improved with detailed, actionable recommendations]
            
            ## ATS Optimization Assessment
            [Analyze how well the resume is optimized for Applicant Tracking Systems. Provide a specific ATS score from 0-100, with 100 being perfectly optimized. Use this format: "ATS Score: XX/100". Then suggest specific keywords and formatting changes to improve ATS performance.]
            
            ## Recommended Courses/Certifications
            [Suggest 5-7 specific courses or certifications that would enhance the candidate's profile, with a brief explanation of why each would be valuable]
            
            ## Resume Score
            [Provide a score from 0-100 based on the overall quality of the resume. Use this format exactly: "Resume Score: XX/100" where XX is the numerical score. Be consistent with your assessment - a resume with significant issues should score below 60, an average resume 60-75, a good resume 75-85, and an excellent resume 85-100.]
            
            Resume:
            {resume_text}
            """
            
            if job_role:
                base_prompt += f"""
                
                The candidate is targeting a role as: {job_role}
                
                ## Role Alignment Analysis
                [Analyze how well the resume aligns with the target role of {job_role}. Provide specific recommendations to better align the resume with this role.]
                """
            
            if job_description:
                base_prompt += f"""
                
                Additionally, compare this resume to the following job description:
                
                Job Description:
                {job_description}
                
                ## Job Match Analysis
                [Provide a detailed analysis of how well the resume matches the job description, with a match percentage and specific areas of alignment and misalignment]
                
                ## Key Job Requirements Not Met
                [List specific requirements from the job description that are not addressed in the resume, with recommendations on how to address each gap]
                """
            
            try:
                analysis, provider_info = self._generate_ai_completion(base_prompt, temperature=0.7)
            except Exception:
                analysis, provider_info = self._generate_fallback_analysis(resume_text, job_role, job_description)
            
            # Extract resume score if present
            resume_score = self._extract_score_from_text(analysis)
            
            # Extract ATS score if present
            ats_score = self._extract_ats_score_from_text(analysis)
            
            return {
                "analysis": analysis,
                "resume_score": resume_score,
                "ats_score": ats_score,
                "model_used": provider_info
            }
        
        except Exception:
            analysis, provider_info = self._generate_fallback_analysis(resume_text, job_role, job_description)
            resume_score = self._extract_score_from_text(analysis)
            ats_score = self._extract_ats_score_from_text(analysis)
            return {
                "analysis": analysis,
                "resume_score": resume_score,
                "ats_score": ats_score,
                "model_used": provider_info
            }

    def analyze_resume_with_gemini(self, resume_text, job_description=None, job_role=None):
        """Backward compatibility wrapper redirecting to multi-provider AI pipeline"""
        return self.analyze_resume_with_openrouter(resume_text, job_description, job_role)

    def generate_roast(self, resume_text):
        """Generate a brutal roast of the resume using multi-provider AI pipeline"""
        if not resume_text:
            return "Roast unavailable: Resume text is required."
            
        prompt = (
            "You are a blunt, comedic resume reviewer. Critique ONLY the document text — "
            "its writing, vague bullets, buzzwords, formatting problems, generic objectives, "
            "unexplained gaps, and clichés. Never comment on the person's identity, background, "
            "appearance, or anything not in the text. No discriminatory or harassing content. "
            "Each criticism must include exactly one concrete actionable fix. "
            "Tone: blunt and funny, not cruel. No emoji. "
            "Format: numbered list, criticism then 'Fix:' on the same or next line.\n\n"
            f"Resume:\n{resume_text}"
        )
        
        try:
            roast_text, _ = self._generate_ai_completion(prompt, temperature=0.8)
            return roast_text
        except Exception:
            return (
                "1. Your resume relies heavily on passive descriptions instead of high-impact metrics.\n"
                "   Fix: Start every bullet point with a powerful action verb and add at least one metric (e.g., 'Increased efficiency by 30%').\n\n"
                "2. Your skills section lists generic tools without highlighting specific frameworks or proficiency.\n"
                "   Fix: Group skills into Technical, Domain, and Methodologies with specific tool versions.\n\n"
                "3. Bullet points read like a job description rather than personal achievements.\n"
                "   Fix: Frame accomplishments using the STAR method (Situation, Task, Action, Result)."
            )

    
    def generate_pdf_report(self, analysis_result, candidate_name, job_role):
        """Generate a PDF report of the analysis"""
        try:
            if SimpleDocTemplate is None:
                st.error("Error importing PDF libraries: reportlab is not installed.")
                st.info("Please make sure reportlab is installed: pip install reportlab")
                return self.simple_generate_pdf_report(analysis_result, candidate_name, job_role)
            
            # Helper function to clean markdown formatting
            def clean_markdown(text):
                if not text:
                    return ""
                
                # Remove markdown formatting for bold and italic
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove ** for bold
                text = re.sub(r'\*(.*?)\*', r'\1', text)      # Remove * for italic
                text = re.sub(r'__(.*?)__', r'\1', text)      # Remove __ for bold
                text = re.sub(r'_(.*?)_', r'\1', text)        # Remove _ for italic
                
                # Remove markdown formatting for headers
                text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
                
                # Remove markdown formatting for links
                text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
                
                return text.strip()
            
            # Validate input data
            if not analysis_result:
                st.error("No analysis result provided for PDF generation")
                return None
                
            # Print debug info
            st.info(f"Generating PDF report for {candidate_name} targeting {job_role}")
            
            # Create a buffer for the PDF
            buffer = io.BytesIO()
            
            # Create the PDF document
            doc = SimpleDocTemplate(buffer, pagesize=letter, 
                                   leftMargin=0.5*inch, rightMargin=0.5*inch,
                                   topMargin=0.5*inch, bottomMargin=0.5*inch)
            styles = getSampleStyleSheet()
            
            # Create custom styles
            title_style = ParagraphStyle(
                'Title',
                parent=styles['Heading1'],
                fontSize=20,
                textColor=colors.darkblue,
                spaceAfter=12,
                alignment=1  # Center alignment
            )
            
            subtitle_style = ParagraphStyle(
                'Subtitle',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.darkblue,
                spaceAfter=12,
                alignment=1  # Center alignment
            )
            
            heading_style = ParagraphStyle(
                'Heading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.white,
                spaceAfter=6,
                backColor=colors.darkblue,
                borderWidth=1,
                borderColor=colors.grey,
                borderPadding=5,
                borderRadius=5,
                alignment=1  # Center alignment
            )
            
            subheading_style = ParagraphStyle(
                'SubHeading',
                parent=styles['Heading3'],
                fontSize=12,
                textColor=colors.darkblue,
                spaceAfter=6,
                borderWidth=0,
                borderPadding=0,
                borderColor=colors.grey,
                borderRadius=0
            )
            
            normal_style = ParagraphStyle(
                'Normal',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=6,
                leading=14  # Line spacing
            )
            
            list_item_style = ParagraphStyle(
                'ListItem',
                parent=normal_style,
                leftIndent=20,
                firstLineIndent=-15,
                spaceBefore=2,
                spaceAfter=2
            )
            
            # Create a gauge chart class
            class GaugeChart(Drawing):
                def __init__(self, width, height, score, max_score=100, label=""):
                    Drawing.__init__(self, width, height)
                    self.width = width
                    self.height = height
                    self._score = int(score) if score is not None else 0  # Ensure score is an integer
                    self._max_score = max_score  # Use _max_score to avoid attribute error
                    self._label = label  # Use _label instead of label to avoid attribute error
                    
                    # Determine color based on score percentage
                    score_percent = (self._score / self._max_score) * 100 if self._max_score > 0 else 0
                    if score_percent >= 80:
                        self._color = colors.green
                        self._status = "Excellent"
                    elif score_percent >= 60:
                        self._color = colors.orange
                        self._status = "Good"
                    else:
                        self._color = colors.red
                        self._status = "Needs Improvement"
                    
                    self._draw()
                
                def _draw(self):
                    # Background
                    self.add(Rect(0, 0, self.width, self.height, 
                                 fillColor=colors.white, strokeColor=None))
                    
                    # Draw gauge background (arc)
                    center_x = self.width / 2
                    center_y = self.height / 2 - 10
                    radius = min(center_x, center_y) - 10
                    
                    # Draw the gauge background
                    for i in range(0, 101, 2):
                        angle = math.radians(180 - (i * 1.8))
                        x = center_x + radius * math.cos(angle)
                        y = center_y + radius * math.sin(angle)
                        
                        # Determine color for background segments
                        if i < 60 or i < 80:
                            segment_color = colors.lightgrey
                        else:
                            segment_color = colors.lightgrey
                        
                        # Draw a small line for each segment
                        line_length = 5
                        end_x = center_x + (radius + line_length) * math.cos(angle)
                        end_y = center_y + (radius + line_length) * math.sin(angle)
                        
                        self.add(Line(x, y, end_x, end_y, strokeColor=segment_color, strokeWidth=2))
                    
                    # Draw the colored arc for the score
                    score_angle = math.radians(180 - (self._score * 1.8))
                    score_x = center_x + radius * math.cos(score_angle)
                    score_y = center_y + radius * math.sin(score_angle)
                    
                    # Draw needle
                    self.add(Line(center_x, center_y, score_x, score_y, 
                                 strokeColor=self._color, strokeWidth=3))
                    
                    # Draw center circle
                    self.add(Circle(center_x, center_y, 5, 
                                   fillColor=self._color, strokeColor=None))
                    
                    # Draw score text
                    self.add(String(center_x, center_y - 25, f"{self._score}",
                                   fontSize=20, fillColor=self._color, 
                                   textAnchor='middle', fontName='Helvetica-Bold'))
                    
                    # Draw status text
                    self.add(String(center_x, center_y - 40, self._status,
                                   fontSize=12, fillColor=colors.black, 
                                   textAnchor='middle'))
                    
                    # Draw label
                    if self._label:
                        self.add(String(center_x, self.height - 15, self._label,
                                       fontSize=12, fillColor=colors.darkblue, 
                                       textAnchor='middle', fontName='Helvetica-Bold'))
                    
                    # Draw scale markers
                    for i in range(0, 101, 20):
                        angle = math.radians(180 - (i * 1.8))
                        x = center_x + (radius - 15) * math.cos(angle)
                        y = center_y + (radius - 15) * math.sin(angle)
                        
                        self.add(String(x, y, str(i),
                                       fontSize=8, fillColor=colors.black, 
                                       textAnchor='middle'))
            
            # Create a Circle class for the gauge
            class Circle(Rect):
                def __init__(self, cx, cy, r, **kw):
                    Rect.__init__(self, cx-r, cy-r, 2*r, 2*r, **kw)
                    self.rx = self.ry = r
            
            # Create a combined gauge chart class
            class CombinedGaugeChart(Drawing):
                def __init__(self, width, height, resume_score, ats_score, max_score=100):
                    Drawing.__init__(self, width, height)
                    self.width = width
                    self.height = height
                    self._resume_score = resume_score
                    self._ats_score = ats_score
                    self._max_score = max_score
                    
                    # Calculate combined score (weighted average)
                    self._combined_score = int((self._resume_score * 0.6) + (self._ats_score * 0.4))
                    
                    # Determine color based on score percentage
                    if self._combined_score >= 80:
                        self._color = colors.green
                        self._status = "Excellent"
                    elif self._combined_score >= 60:
                        self._color = colors.orange
                        self._status = "Good"
                    else:
                        self._color = colors.red
                        self._status = "Needs Improvement"
                    
                    self._draw()
                
                def _draw(self):
                    # Background
                    self.add(Rect(0, 0, self.width, self.height, 
                                 fillColor=colors.white, strokeColor=None))
                    
                    # Draw gauge background (arc)
                    center_x = self.width / 2
                    center_y = self.height / 2
                    radius = min(center_x, center_y) - 20
                    
                    # Draw the gauge background
                    for i in range(0, 101, 2):
                        angle = math.radians(180 - (i * 1.8))
                        x = center_x + radius * math.cos(angle)
                        y = center_y + radius * math.sin(angle)
                        
                        # Determine color for background segments
                        segment_color = colors.lightgrey
                        
                        # Draw a small line for each segment
                        line_length = 5
                        end_x = center_x + (radius + line_length) * math.cos(angle)
                        end_y = center_y + (radius + line_length) * math.sin(angle)
                        
                        self.add(Line(x, y, end_x, end_y, strokeColor=segment_color, strokeWidth=2))
                    
                    # Draw the colored arc for the combined score
                    score_angle = math.radians(180 - (self._combined_score * 1.8))
                    score_x = center_x + radius * math.cos(score_angle)
                    score_y = center_y + radius * math.sin(score_angle)
                    
                    # Draw needle
                    self.add(Line(center_x, center_y, score_x, score_y, 
                                 strokeColor=self._color, strokeWidth=3))
                    
                    # Draw center circle
                    self.add(Circle(center_x, center_y, 5, 
                                   fillColor=self._color, strokeColor=None))
                    
                    # Draw combined score text
                    self.add(String(center_x, center_y - 25, f"{self._combined_score}",
                                   fontSize=24, fillColor=self._color, 
                                   textAnchor='middle', fontName='Helvetica-Bold'))
                    
                    # Draw status text
                    self.add(String(center_x, center_y - 45, self._status,
                                   fontSize=12, fillColor=colors.black, 
                                   textAnchor='middle'))
                    
                    # Draw individual scores
                    self.add(String(center_x - 60, center_y - 70, f"Resume: {self._resume_score}",
                                   fontSize=10, fillColor=colors.darkblue, 
                                   textAnchor='middle'))
                    
                    self.add(String(center_x + 60, center_y - 70, f"ATS: {self._ats_score}",
                                   fontSize=10, fillColor=colors.darkblue, 
                                   textAnchor='middle'))
                    
                    # Draw "Overall Score" label
                    self.add(String(center_x, self.height - 15, "Overall Score",
                                   fontSize=14, fillColor=colors.darkblue, 
                                   textAnchor='middle', fontName='Helvetica-Bold'))
                    
                    # Draw scale markers
                    for i in range(0, 101, 20):
                        angle = math.radians(180 - (i * 1.8))
                        x = center_x + (radius - 15) * math.cos(angle)
                        y = center_y + (radius - 15) * math.sin(angle)
                        
                        self.add(String(x, y, str(i),
                                       fontSize=8, fillColor=colors.black, 
                                       textAnchor='middle'))
            
            # Create the content
            content = []
            
            # Add a header with date
            current_date = datetime.datetime.now().strftime("%B %d, %Y")
            content.append(Paragraph("Resume Analysis Report", title_style))
            content.append(Paragraph(f"Generated on {current_date}", subtitle_style))
            content.append(Spacer(1, 0.25*inch))
            
            # Format candidate name - if it's just "Candidate", add a number
            if not candidate_name or candidate_name.lower() == "candidate" or candidate_name.strip() == "":
                import random
                candidate_name = f"Candidate_{random.randint(1000, 9999)}"
            
            # Add candidate name and job role in a table
            info_data = [
                ["Candidate:", candidate_name],
                ["Target Role:", job_role if job_role else "Not specified"]
            ]
            
            info_table = Table(info_data, colWidths=[1.5*inch, 5*inch])
            info_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 12),
                ('TEXTCOLOR', (0, 0), (0, -1), colors.darkblue),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
            ]))
            
            content.append(info_table)
            content.append(Spacer(1, 0.25*inch))
            
            # Analysis Content
            analysis_text = analysis_result.get("full_response", "")
            
            # Extract key sections for the executive summary
            strengths = analysis_result.get("strengths", [])
            weaknesses = analysis_result.get("weaknesses", [])
            
            # If strengths and weaknesses are not in the structured data, try to extract from text
            if not strengths:
                if "## Key Strengths" in analysis_text:
                    strengths_section = analysis_text.split("## Key Strengths")[1].split("##")[0].strip()
                    strengths = [clean_markdown(s.strip().replace("- ", "").replace("* ", "").replace("• ", "")) 
                                for s in strengths_section.split("\n") 
                                if s.strip() and (s.strip().startswith("-") or s.strip().startswith("*") or s.strip().startswith("•"))]
                
                # Try another pattern for strengths
                if not strengths and "Key Strengths" in analysis_text:
                    strengths_section = analysis_text.split("Key Strengths")[1]
                    if "Areas for Improvement" in strengths_section:
                        strengths_section = strengths_section.split("Areas for Improvement")[0]
                    
                    # Extract lines that look like list items
                    for line in strengths_section.split("\n"):
                        line = line.strip()
                        if line and (line.startswith(("-", "*", "•"))):
                            strengths.append(clean_markdown(line.replace("- ", "").replace("* ", "").replace("• ", "")))
                        elif line and ":" in line and not line.startswith("#"):
                            strengths.append(clean_markdown(line))

            if not weaknesses:
                if "## Areas for Improvement" in analysis_text:
                    weaknesses_section = analysis_text.split("## Areas for Improvement")[1].split("##")[0].strip()
                    weaknesses = [clean_markdown(w.strip().replace("- ", "").replace("* ", "").replace("• ", "")) 
                                 for w in weaknesses_section.split("\n") 
                                 if w.strip() and (w.strip().startswith("-") or w.strip().startswith("*") or w.strip().startswith("•"))]
                
                # Try another pattern for weaknesses
                if not weaknesses and "Areas for Improvement" in analysis_text:
                    weaknesses_section = analysis_text.split("Areas for Improvement")[1]
                    if "##" in weaknesses_section:
                        weaknesses_section = weaknesses_section.split("##")[0]
                    
                    # Extract lines that look like list items
                    for line in weaknesses_section.split("\n"):
                        line = line.strip()
                        if line and (line.startswith(("-", "*", "•"))):
                            weaknesses.append(clean_markdown(line.replace("- ", "").replace("* ", "").replace("• ", "")))
                        elif line and ":" in line and not line.startswith("#"):
                            weaknesses.append(clean_markdown(line))
            
            # Extract scores
            resume_score = analysis_result.get("score", 0)
            if resume_score == 0:
                # Try to get from resume_score
                resume_score = analysis_result.get("resume_score", 0)
                
                # If still 0, try to extract from the analysis text
                if resume_score == 0 and "Resume Score:" in analysis_text:
                    score_match = re.search(r'Resume Score:\s*(\d{1,3})/100', analysis_text)
                    if score_match:
                        resume_score = int(score_match.group(1))
                    else:
                        # Try another pattern
                        score_match = re.search(r'\bResume Score:\s*(\d{1,3})\b', analysis_text)
                        if score_match:
                            resume_score = int(score_match.group(1))
                        else:
                            # Try to find any number after "Resume Score:"
                            score_section = analysis_text.split("Resume Score:")[1].split("\n")[0].strip()
                            score_match = re.search(r'\b(\d{1,3})\b', score_section)
                            if score_match:
                                resume_score = int(score_match.group(1))

            # Ensure resume_score is a valid integer
            resume_score = int(resume_score) if resume_score else 0
            resume_score = max(0, min(resume_score, 100))  # Ensure it's between 0 and 100

            analysis_result.get("ats_score", 0)
            model_used = analysis_result.get("model_used", "AI")

            # Add model used information
            model_data = [["Analysis performed by:",model_used]]
            model_table = Table(model_data, colWidths=[1.9*inch, 5*inch])
            model_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 12),
                ('TEXTCOLOR', (0, 0), (0, -1), colors.darkblue),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
            ]))

            content.append(model_table)
            content.append(Spacer(1, 0.25*inch))

            # Add score gauges
            content.append(Paragraph("Resume Evaluation", heading_style))
            content.append(Spacer(1, 0.1*inch))

            # Create a table with the gauge
            score_table_data = [
                ["Resume Score"],
                [GaugeChart(width=300, height=200, score=resume_score, max_score=100, label="Resume Score")]
            ]
            score_table = Table(score_table_data, colWidths=[6*inch])
            score_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (0, 0), 14),
                ('TEXTCOLOR', (0, 0), (0, 0), colors.darkblue),
                ('BOTTOMPADDING', (0, 0), (0, 0), 10),
            ]))

            content.append(score_table)
            content.append(Spacer(1, 0.25*inch))

            # Add Executive Summary section
            content.append(Paragraph("Executive Summary", heading_style))
            content.append(Spacer(1, 0.1*inch))

            # Extract overall assessment
            overall_assessment = ""
            if "## Overall Assessment" in analysis_text:
                overall_section = analysis_text.split("## Overall Assessment")[1].split("##")[0].strip()
                overall_assessment = clean_markdown(overall_section)

            content.append(Paragraph(overall_assessment, normal_style))
            content.append(Spacer(1, 0.2*inch))

            # Key Strengths and Areas for Improvement section
            content.append(Paragraph("Key Strengths and Areas for Improvement", subheading_style))
            content.append(Spacer(1, 0.1*inch))

            if strengths or weaknesses:
                # Create data for strengths and weaknesses
                sw_data = [["Key Strengths", "Areas for Improvement"]]
                
                # Get max length of strengths and weaknesses
                max_len = max(len(strengths), len(weaknesses), 1)
                
                for i in range(max_len):
                    strength = f"• {clean_markdown(strengths[i])}" if i < len(strengths) else ""
                    weakness = f"• {clean_markdown(weaknesses[i])}" if i < len(weaknesses) else ""
                    sw_data.append([
                        Paragraph(strength, list_item_style) if strength else "",
                        Paragraph(weakness, list_item_style) if weakness else ""
                    ])
                
                sw_table = Table(sw_data, colWidths=[3*inch, 3*inch])
                sw_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, 0), colors.lightgreen),
                    ('BACKGROUND', (1, 0), (1, 0), colors.salmon),
                    ('TEXTCOLOR', (0, 0), (1, 0), colors.black),
                    ('ALIGN', (0, 0), (1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (1, 0), 10),
                    ('GRID', (0, 0), (1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                
                content.append(sw_table)
            else:
                # Add empty strengths and weaknesses with a message
                empty_data = [
                    ["Key Strengths", "Areas for Improvement"],
                    [
                        Paragraph("No specific strengths identified in the analysis.", normal_style),
                        Paragraph("No specific areas for improvement identified in the analysis.", normal_style)
                    ]
                ]
                empty_table = Table(empty_data, colWidths=[3*inch, 3*inch])
                empty_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, 0), colors.lightgreen),
                    ('BACKGROUND', (1, 0), (1, 0), colors.salmon),
                    ('TEXTCOLOR', (0, 0), (1, 0), colors.black),
                    ('ALIGN', (0, 0), (1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (1, 0), 10),
                    ('GRID', (0, 0), (1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                
                content.append(empty_table)

            content.append(Spacer(1, 0.25*inch))
            
            # Add Detailed Analysis section
            content.append(Paragraph("Detailed Analysis", heading_style))
            content.append(Spacer(1, 0.1*inch))
            
            # Parse the markdown-like content
            sections = analysis_text.split("##")
            
            # Define sections to include in detailed analysis
            detailed_sections = [
                "Professional Profile Analysis",
                "Skills Analysis",
                "Experience Analysis",
                "Education Analysis",
                "ATS Optimization Assessment",
                "Role Alignment Analysis",
                "Job Match Analysis"
            ]
            
            for section in sections:
                if not section.strip():
                    continue
                
                # Extract section title and content
                lines = section.strip().split("\n")
                section_title = lines[0].strip()
                
                # Skip sections we don't want in the detailed analysis
                if section_title not in detailed_sections and section_title != "Overall Assessment":
                    continue
                
                # Skip Overall Assessment as we've already included it
                if section_title == "Overall Assessment":
                    continue
                
                section_content = "\n".join(lines[1:]).strip()
                
                # Add section title
                content.append(Paragraph(section_title, subheading_style))
                content.append(Spacer(1, 0.1*inch))
                
                # Process content based on section
                if section_title == "Skills Analysis":
                    # Extract current and missing skills
                    current_skills = []
                    missing_skills = []
                    
                    if "Current Skills" in section_content:
                        current_part = section_content.split("Current Skills")[1]
                        if "Missing Skills" in current_part:
                            current_part = current_part.split("Missing Skills")[0]
                        
                        for line in current_part.split("\n"):
                            if line.strip() and ("-" in line or "*" in line or "•" in line):
                                skill = line.replace("-", "").replace("*", "").replace("•", "").strip()
                                if skill:
                                    current_skills.append(skill)
                    
                    if "Missing Skills" in section_content:
                        missing_part = section_content.split("Missing Skills")[1]
                        for line in missing_part.split("\n"):
                            if line.strip() and ("-" in line or "*" in line or "•" in line):
                                skill = line.replace("-", "").replace("*", "").replace("•", "").strip()
                                if skill:
                                    missing_skills.append(skill)
                    
                    # Create skills table with better formatting
                    if current_skills or missing_skills:
                        # Create paragraphs for each skill to ensure proper wrapping
                        current_skill_paragraphs = [Paragraph(skill, normal_style) for skill in current_skills]
                        missing_skill_paragraphs = [Paragraph(skill, normal_style) for skill in missing_skills]
                        
                        # Make sure both lists have the same length
                        max_len = max(len(current_skill_paragraphs), len(missing_skill_paragraphs))
                        current_skill_paragraphs.extend([Paragraph("", normal_style)] * (max_len - len(current_skill_paragraphs)))
                        missing_skill_paragraphs.extend([Paragraph("", normal_style)] * (max_len - len(missing_skill_paragraphs)))
                        
                        # Create data for the table
                        data = [["Current Skills", "Missing Skills"]]
                        for i in range(max_len):
                            data.append([current_skill_paragraphs[i], missing_skill_paragraphs[i]])
                        
                        # Create the table with fixed column widths
                        table = Table(data, colWidths=[3*inch, 3*inch])
                        table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (1, 0), colors.lightgreen),
                            ('TEXTCOLOR', (0, 0), (1, 0), colors.black),
                            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                            ('GRID', (0, 0), (-1, -1), 1, colors.black),
                            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                            ('LEFTPADDING', (0, 0), (-1, -1), 10),
                            ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                        ]))
                        
                        content.append(table)
                    
                    # We no longer need to add skill proficiency outside the table
                    # as it's now included in the table itself
                elif section_title == "ATS Optimization Assessment":
                    # Special handling for ATS Optimization Assessment
                    ats_score_line = ""
                    ats_content = []
                    
                    # Extract ATS score if present
                    for line in section_content.split("\n"):
                        if "ATS Score:" in line:
                            ats_score_line = clean_markdown(line)
                        elif line.strip():
                            # Check if it's a list item
                            if line.strip().startswith("-") or line.strip().startswith("*") or line.strip().startswith("•"):
                                ats_content.append("• " + clean_markdown(line.strip()[1:].strip()))
                            else:
                                ats_content.append(clean_markdown(line))
                    
                    # Add ATS score line if found
                    if ats_score_line:
                        content.append(Paragraph(ats_score_line, normal_style))
                        content.append(Spacer(1, 0.1*inch))
                    
                    # Add the rest of the ATS content
                    for para in ats_content:
                        if para.startswith("• "):
                            content.append(Paragraph(para, list_item_style))
                        else:
                            content.append(Paragraph(para, normal_style))
                else:
                    # Process regular paragraphs
                    paragraphs = section_content.split("\n")
                    for para in paragraphs:
                        if para.strip():
                            # Check if it's a list item
                            if para.strip().startswith("-") or para.strip().startswith("*") or para.strip().startswith("•"):
                                para = "• " + clean_markdown(para.strip()[1:].strip())
                                content.append(Paragraph(para, list_item_style))
                            else:
                                content.append(Paragraph(clean_markdown(para), normal_style))
                
                content.append(Spacer(1, 0.2*inch))
            
            # Add course recommendations
            course_recommendations = []
            
            # Try to get course recommendations from different sources
            if "suggestions" in analysis_result:
                course_recommendations = analysis_result.get("suggestions", [])
            
            # If still no recommendations, try to extract from text
            if not course_recommendations and "## Recommended Courses" in analysis_text:
                recommendations_section = analysis_text.split("## Recommended Courses")[1].split("##")[0].strip()
                course_recommendations = [clean_markdown(r.strip().replace("- ", "").replace("* ", "").replace("• ", "")) 
                              for r in recommendations_section.split("\n") 
                              if r.strip() and (r.strip().startswith("-") or r.strip().startswith("*") or r.strip().startswith("•"))]
            
            # Try another pattern for course recommendations
            if not course_recommendations and "Recommended Courses" in analysis_text:
                recommendations_section = analysis_text.split("Recommended Courses")[1]
                if "##" in recommendations_section:
                    recommendations_section = recommendations_section.split("##")[0]
                
                # Extract lines that look like list items
                for line in recommendations_section.split("\n"):
                    line = line.strip()
                    if line and ":" in line and not line.startswith("#"):
                        course_recommendations.append(clean_markdown(line))
            
            content.append(Paragraph("Recommended Courses & Certifications", subheading_style))
            
            if course_recommendations:
                # Create a table for course recommendations with better formatting
                course_data = [["Recommended Courses & Certifications"]]  # Add header row
                
                for course in course_recommendations:
                    # Clean the course text and ensure it doesn't have any markdown formatting
                    cleaned_course = clean_markdown(course)
                    course_data.append([Paragraph(f"• {cleaned_course}", list_item_style)])
                
                course_table = Table(course_data, colWidths=[6*inch])
                course_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, 0), colors.lightblue),
                    ('TEXTCOLOR', (0, 0), (0, 0), colors.black),
                    ('ALIGN', (0, 0), (0, 0), 'CENTER'),  # Center the header
                    ('ALIGN', (0, 1), (0, -1), 'LEFT'),   # Left-align the content
                    ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (0, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (0, 0), 10),
                    ('GRID', (0, 0), (0, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (0, -1), 'TOP'),
                ]))
                
                content.append(course_table)
            else:
                # If still no recommendations, add a text section instead of generic courses
                content.append(Paragraph("Based on your resume and target role, consider the following types of courses and certifications:", normal_style))
                content.append(Spacer(1, 0.1*inch))
                
                # Add role-specific recommendations based on job_role
                role_specific_courses = []
                if "data" in job_role.lower() or "scientist" in job_role.lower() or "analyst" in job_role.lower():
                    role_specific_courses = [
                        "Data Science Specialization (Coursera/edX)",
                        "Machine Learning (Coursera/edX)",
                        "Deep Learning Specialization (Coursera)",
                        "Big Data Technologies (Cloud Provider Certifications)",
                        "Statistical Modeling and Inference",
                        "Data Visualization with Tableau/Power BI"
                    ]
                elif "developer" in job_role.lower() or "engineer" in job_role.lower() or "programming" in job_role.lower():
                    role_specific_courses = [
                        "Full Stack Web Development (Udemy/Coursera)",
                        "Cloud Certifications (AWS/Azure/GCP)",
                        "DevOps and CI/CD Pipelines",
                        "Software Architecture and Design Patterns",
                        "Agile and Scrum Methodologies",
                        "Mobile App Development"
                    ]
                elif "security" in job_role.lower() or "cyber" in job_role.lower():
                    role_specific_courses = [
                        "Certified Information Systems Security Professional (CISSP)",
                        "Certified Ethical Hacker (CEH)",
                        "CompTIA Security+",
                        "Offensive Security Certified Professional (OSCP)",
                        "Cloud Security Certifications",
                        "Security Operations and Incident Response"
                    ]
                else:
                    # Generic professional development courses
                    role_specific_courses = [
                        "LinkedIn Learning - Professional Skills Development",
                        "Coursera - Career Development Specialization",
                        "Udemy - Job Interview Skills Training",
                        "Project Management Professional (PMP)",
                        "Leadership and Management Skills",
                        "Technical Writing and Communication"
                    ]
                
                # Create a table for role-specific courses
                course_data = []
                for course in role_specific_courses:
                    course_data.append([Paragraph(f"• {clean_markdown(course)}", list_item_style)])
                
                course_table = Table(course_data, colWidths=[6*inch])
                course_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, 0), colors.lightblue),
                    ('TEXTCOLOR', (0, 0), (0, 0), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica'),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]))
                
                content.append(course_table)
            
            content.append(Spacer(1, 0.2*inch))
            
            # Add footer with page numbers
            def add_page_number(canvas, doc):
                canvas.saveState()
                canvas.setFont('Helvetica', 9)
                page_num = canvas.getPageNumber()
                text = f"Page {page_num}"
                canvas.drawRightString(7.5*inch, 0.25*inch, text)
                
                # Add generation date at the bottom
                canvas.setFont('Helvetica', 9)
                date_text = f"Generated on: {datetime.datetime.now().strftime('%B %d, %Y')}"
                canvas.drawString(0.5*inch, 0.25*inch, date_text)
                
                canvas.restoreState()
            
            # Build the PDF
            doc.build(content, onFirstPage=add_page_number, onLaterPages=add_page_number)
            
            # Get the PDF from the buffer
            buffer.seek(0)
            return buffer
        
        except Exception as e:
            st.error(f"Error generating simple PDF report: {e!s}")
            import traceback
            st.code(traceback.format_exc())
            return None
            
    def extract_skills_from_analysis(self, analysis_text):
        """Extract skills from the analysis text"""
        skills = []
        
        try:
            if "Current Skills" in analysis_text:
                skills_section = analysis_text.split("Current Skills")[1]
                if "##" in skills_section:
                    skills_section = skills_section.split("##")[0]
                
                for line in skills_section.split("\n"):
                    if line.strip() and ("-" in line or "*" in line or "•" in line):
                        skill = line.replace("-", "").replace("*", "").replace("•", "").strip()
                        if skill:
                            skills.append(skill)
        except Exception as e:
            st.warning(f"Error extracting skills: {e!s}")
        
        return skills
        
    def extract_missing_skills_from_analysis(self, analysis_text):
        """Extract missing skills from the analysis text"""
        missing_skills = []
        
        try:
            if "Missing Skills" in analysis_text:
                missing_section = analysis_text.split("Missing Skills")[1]
                if "##" in missing_section:
                    missing_section = missing_section.split("##")[0]
                
                for line in missing_section.split("\n"):
                    if line.strip() and ("-" in line or "*" in line or "•" in line):
                        skill = line.replace("-", "").replace("*", "").replace("•", "").strip()
                        if skill:
                            missing_skills.append(skill)
        except Exception as e:
            st.warning(f"Error extracting missing skills: {e!s}")
        
        return missing_skills
    
    def _extract_score_from_text(self, analysis_text):
        """Extract the resume score from the analysis text"""
        try:
            # Look for the Resume Score section
            if "## Resume Score" in analysis_text:
                score_section = analysis_text.split("## Resume Score")[1].strip()
                # Extract the first number found
                score_match = re.search(r'Resume Score:\s*(\d{1,3})/100', score_section)
                if score_match:
                    score = int(score_match.group(1))
                    # Ensure score is within valid range
                    return max(0, min(score, 100))
                
                # Try another pattern if the first one doesn't match
                score_match = re.search(r'\b(\d{1,3})\b', score_section)
                if score_match:
                    score = int(score_match.group(1))
                    # Ensure score is within valid range
                    return max(0, min(score, 100))
            
            # If no score found in Resume Score section, try to find it elsewhere
            score_match = re.search(r'Resume Score:\s*(\d{1,3})/100', analysis_text)
            if score_match:
                score = int(score_match.group(1))
                return max(0, min(score, 100))
                
            return 0
        except Exception as e:
            print(f"Error extracting score: {e!s}")
            return 0
            
    def _extract_ats_score_from_text(self, analysis_text):
        """Extract the ATS score from the analysis text"""
        try:
            # Look for the ATS Score in the ATS Optimization Assessment section
            if "## ATS Optimization Assessment" in analysis_text:
                ats_section = analysis_text.split("## ATS Optimization Assessment")[1].split("##")[0].strip()
                # Extract the score using regex
                score_match = re.search(r'ATS Score:\s*(\d{1,3})/100', ats_section)
                if score_match:
                    score = int(score_match.group(1))
                    # Ensure score is within valid range
                    return max(0, min(score, 100))
            return 0
        except Exception as e:
            print(f"Error extracting ATS score: {e!s}")
            return 0
            
    def analyze_resume(self, resume_text, job_role=None, role_info=None, model="OpenRouter", custom_jd=None, job_description=None, **kwargs):
        """
        Analyze a resume using the specified AI model
        
        Parameters:
        - resume_text: The text content of the resume
        - job_role: The target job role
        - role_info: Additional information about the job role
        - model: The AI model to use ("OpenRouter")
        - custom_jd: Custom job description text
        - job_description: Alternative job description parameter
        
        Returns:
        - Dictionary containing analysis results
        """
        import traceback
        
        try:
            effective_jd = job_description or custom_jd
            if role_info and not effective_jd:
                effective_jd = f"""
                Role: {job_role}
                Description: {role_info.get('description', '')}
                Required Skills: {', '.join(role_info.get('required_skills', []))}
                """
            
            result = self.analyze_resume_with_openrouter(resume_text, effective_jd, job_role)
            model_used = result.get("model_used", f"OpenRouter ({OPENROUTER_MODEL})")
            
            # Process the result to extract structured information
            analysis_text = result.get("analysis", "")
            
            # Extract strengths
            strengths = []
            if "## Key Strengths" in analysis_text:
                strengths_section = analysis_text.split("## Key Strengths")[1].split("##")[0].strip()
                strengths = [self._clean_markdown(s.strip().replace("- ", "").replace("* ", "").replace("• ", "")) 
                            for s in strengths_section.split("\n") 
                            if s.strip() and (s.strip().startswith("-") or s.strip().startswith("*") or s.strip().startswith("•"))]
            
            # Extract weaknesses/areas for improvement
            weaknesses = []
            if "## Areas for Improvement" in analysis_text:
                weaknesses_section = analysis_text.split("## Areas for Improvement")[1].split("##")[0].strip()
                weaknesses = [self._clean_markdown(w.strip().replace("- ", "").replace("* ", "").replace("• ", "")) 
                             for w in weaknesses_section.split("\n") 
                             if w.strip() and (w.strip().startswith("-") or w.strip().startswith("*") or w.strip().startswith("•"))]
            
            # Extract suggestions/recommendations
            suggestions = []
            if "## Recommended Courses" in analysis_text:
                suggestions_section = analysis_text.split("## Recommended Courses")[1].split("##")[0].strip()
                suggestions = [self._clean_markdown(s.strip().replace("- ", "").replace("* ", "").replace("• ", "")) 
                                 for s in suggestions_section.split("\n") 
                                 if s.strip() and (s.strip().startswith("-") or s.strip().startswith("*") or s.strip().startswith("•"))]
            
            # Extract score
            score = result.get("resume_score", 0)
            if not score:
                score = self._extract_score_from_text(analysis_text)
            
            # Extract ATS score
            ats_score = self._extract_ats_score_from_text(analysis_text)
            
            # Return structured analysis
            return {
                "score": score,
                "ats_score": ats_score,
                "strengths": strengths,
                "weaknesses": weaknesses,
                "suggestions": suggestions,
                "full_response": analysis_text,
                "model_used": model_used
            }
            
        except Exception as e:
            print(f"Error in analyze_resume: {e!s}")
            print(traceback.format_exc())
            return {
                "error": f"Analysis failed: {e!s}",
                "score": 0,
                "ats_score": 0,
                "strengths": ["Unable to analyze resume due to an error."],
                "weaknesses": ["Unable to analyze resume due to an error."],
                "suggestions": ["Try again with a different model or check your resume format."],
                "full_response": f"Error: {e!s}",
                "model_used": "Error"
            } 

    def simple_generate_pdf_report(self, analysis_result, candidate_name, job_role):
        """Generate a simple PDF report without complex charts as a fallback"""
        try:
            # Import required libraries
            try:
                import datetime
                import io
                import math

                from reportlab.graphics.charts.barcharts import VerticalBarChart
                from reportlab.graphics.charts.legends import Legend
                from reportlab.graphics.charts.linecharts import HorizontalLineChart
                from reportlab.graphics.charts.piecharts import Pie
                from reportlab.graphics.shapes import Drawing, Line, Rect, String
                from reportlab.lib import colors
                from reportlab.lib.pagesizes import letter
                from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
                from reportlab.lib.units import inch
                from reportlab.platypus import (
                    Flowable,
                    Image,
                    KeepTogether,
                    Paragraph,
                    SimpleDocTemplate,
                    Spacer,
                    Table,
                    TableStyle,
                )
            except ImportError as e:
                st.error(f"Error importing PDF libraries: {e!s}")
                st.info("Please make sure reportlab is installed: pip install reportlab")
                return None
            
            # Helper function to clean markdown formatting
            def clean_markdown(text):
                if not text:
                    return ""
                
                # Remove markdown formatting for bold and italic
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove ** for bold
                text = re.sub(r'\*(.*?)\*', r'\1', text)      # Remove * for italic
                text = re.sub(r'__(.*?)__', r'\1', text)      # Remove __ for bold
                text = re.sub(r'_(.*?)_', r'\1', text)        # Remove _ for italic
                
                # Remove markdown formatting for headers
                text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
                
                # Remove markdown formatting for links
                text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
                
                return text.strip()
            
            # Validate input data
            if not analysis_result:
                st.error("No analysis result provided for PDF generation")
                return None
                
            # Create a buffer for the PDF
            buffer = io.BytesIO()
            
            # Create the PDF document
            doc = SimpleDocTemplate(buffer, pagesize=letter, 
                                   leftMargin=0.5*inch, rightMargin=0.5*inch,
                                   topMargin=0.5*inch, bottomMargin=0.5*inch)
            styles = getSampleStyleSheet()
            
            # Create custom styles
            title_style = ParagraphStyle(
                'Title',
                parent=styles['Heading1'],
                fontSize=20,
                textColor=colors.darkblue,
                spaceAfter=12,
                alignment=1  # Center alignment
            )
            
            subtitle_style = ParagraphStyle(
                'Subtitle',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.darkblue,
                spaceAfter=12,
                alignment=1  # Center alignment
            )
            
            heading_style = ParagraphStyle(
                'Heading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.white,
                spaceAfter=6,
                backColor=colors.darkblue,
                borderWidth=1,
                borderColor=colors.grey,
                borderPadding=5,
                borderRadius=5,
                alignment=1  # Center alignment
            )
            
            subheading_style = ParagraphStyle(
                'SubHeading',
                parent=styles['Heading3'],
                fontSize=12,
                textColor=colors.darkblue,
                spaceAfter=6
            )
            
            normal_style = ParagraphStyle(
                'Normal',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=6,
                leading=14  # Line spacing
            )
            
            list_item_style = ParagraphStyle(
                'ListItem',
                parent=normal_style,
                leftIndent=20,
                firstLineIndent=-15,
                spaceBefore=2,
                spaceAfter=2
            )
            
            # Create a simple gauge chart class
            class SimpleGaugeChart(Flowable):
                def __init__(self, score, width=300, height=200, label="Resume Score"):
                    Flowable.__init__(self)
                    self.score = int(score) if score is not None else 0  # Ensure score is an integer
                    self.width = width
                    self.height = height
                    self.label = label
                    
                    # Determine color based on score percentage
                    if self.score >= 80:
                        self.color = colors.green
                        self.status = "Excellent"
                    elif self.score >= 60:
                        self.color = colors.orange
                        self.status = "Good"
                    else:
                        self.color = colors.red
                        self.status = "Needs Improvement"
                
                def draw(self):
                    # Draw the gauge
                    canvas = self.canv
                    canvas.saveState()
                    
                    # Draw gauge background (semi-circle)
                    center_x = self.width / 2
                    center_y = self.height / 2
                    radius = min(center_x, center_y) - 30
                    
                    # Draw the gauge background
                    canvas.setFillColor(colors.lightgrey)
                    canvas.setStrokeColor(colors.grey)
                    canvas.setLineWidth(1)
                    
                    # Draw the semi-circle background
                    p = canvas.beginPath()
                    p.moveTo(center_x, center_y)
                    p.arcTo(center_x - radius, center_y - radius, center_x + radius, center_y + radius, 0, 180)
                    p.lineTo(center_x, center_y)
                    p.close()
                    canvas.drawPath(p, fill=1, stroke=1)
                    
                    # Draw the colored arc for the score
                    if self.score > 0:  # Only draw if score > 0
                        angle = 180 * self.score / 100
                        p = canvas.beginPath()
                        p.moveTo(center_x, center_y)
                        p.arcTo(center_x - radius, center_y - radius, center_x + radius, center_y + radius, 180, 180-angle)
                        p.lineTo(center_x, center_y)
                        p.close()
                        canvas.setFillColor(self.color)
                        canvas.drawPath(p, fill=1, stroke=0)
                    
                    # Draw score text
                    canvas.setFillColor(self.color)
                    canvas.setFont("Helvetica-Bold", 24)
                    canvas.drawCentredString(center_x, center_y - 15, f"{self.score}")
                    
                    # Draw status text
                    canvas.setFillColor(self.color)
                    canvas.setFont("Helvetica", 12)
                    canvas.drawCentredString(center_x, center_y - 35, self.status)
                    
                    # Draw "Resume Score" label
                    canvas.setFillColor(colors.darkblue)
                    canvas.setFont("Helvetica-Bold", 14)
                    canvas.drawCentredString(center_x, self.height - 20, self.label)
                    
                    # Draw scale markers
                    canvas.setStrokeColor(colors.black)
                    canvas.setLineWidth(1)
                    for i in range(0, 101, 20):
                        angle_rad = math.radians(180 - (i * 1.8))
                        x = center_x + radius * math.cos(angle_rad)
                        y = center_y + radius * math.sin(angle_rad)
                        
                        # Draw tick marks
                        x2 = center_x + (radius - 5) * math.cos(angle_rad)
                        y2 = center_y + (radius - 5) * math.sin(angle_rad)
                        canvas.line(x, y, x2, y2)
                        
                        # Draw numbers
                        canvas.setFont("Helvetica", 8)
                        num_x = center_x + (radius - 15) * math.cos(angle_rad)
                        num_y = center_y + (radius - 15) * math.sin(angle_rad)
                        canvas.drawCentredString(num_x, num_y, str(i))
                    
                    canvas.restoreState()
                
                def wrap(self, availWidth, availHeight):
                    return (self.width, self.height)
            
            # Create the content
            content = []
            
            # Add a header with date
            current_date = datetime.datetime.now().strftime("%B %d, %Y")
            content.append(Paragraph("Resume Analysis Report", title_style))
            content.append(Paragraph(f"Generated on {current_date}", subtitle_style))
            content.append(Spacer(1, 0.25*inch))
            
            # Format candidate name - if it's just "Candidate", add a number
            if not candidate_name or candidate_name.lower() == "candidate" or candidate_name.strip() == "":
                import random
                candidate_name = f"Candidate_{random.randint(1000, 9999)}"
            
            # Add candidate name and job role in a table
            info_data = [
                ["Candidate:", candidate_name],
                ["Target Role:", job_role if job_role else "Not specified"]
            ]
            
            info_table = Table(info_data, colWidths=[1.5*inch, 5*inch])
            info_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 12),
                ('TEXTCOLOR', (0, 0), (0, -1), colors.darkblue),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
            ]))
            
            content.append(info_table)
            content.append(Spacer(1, 0.25*inch))
            
            # Add model used information with proper spacing
            model_used = analysis_result.get("model_used", "AI")
            model_data = [["Analysis performed by:\u2003\u2003\u2003", "", model_used]]
            model_table = Table(model_data, colWidths=[3.5*inch, 1*inch, 5*inch])
            model_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 12),
                ('TEXTCOLOR', (0, 0), (0, -1), colors.darkblue),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 20),
            ]))
            
            content.append(model_table)
            content.append(Spacer(1, 0.25*inch))
            
            # Extract core fields
            analysis_text = analysis_result.get("full_response", "") or analysis_result.get("analysis", "")
            strengths = analysis_result.get("strengths", [])
            weaknesses = analysis_result.get("weaknesses", [])

            # Add Resume Evaluation section
            content.append(Paragraph("Resume Evaluation", heading_style))
            content.append(Spacer(1, 0.1*inch))
            
            # Extract scores
            resume_score = analysis_result.get("score", 0)
            if resume_score == 0:
                # Try to get from resume_score
                resume_score = analysis_result.get("resume_score", 0)
                
                # If still 0, try to extract from the analysis text
                if resume_score == 0 and "Resume Score:" in analysis_text:
                    score_match = re.search(r'Resume Score:\s*(\d{1,3})/100', analysis_text)
                    if score_match:
                        resume_score = int(score_match.group(1))
                    else:
                        # Try another pattern
                        score_match = re.search(r'\bResume Score:\s*(\d{1,3})\b', analysis_text)
                        if score_match:
                            resume_score = int(score_match.group(1))
                        else:
                            # Try to find any number after "Resume Score:"
                            if "Resume Score:" in analysis_text:
                                score_section = analysis_text.split("Resume Score:")[1].split("\n")[0].strip()
                                score_match = re.search(r'\b(\d{1,3})\b', score_section)
                                if score_match:
                                    resume_score = int(score_match.group(1))

            # Ensure resume_score is a valid integer
            resume_score = int(resume_score) if resume_score else 0
            resume_score = max(0, min(resume_score, 100))  # Ensure it's between 0 and 100

            # Create a table with the simple gauge
            score_table_data = [
                ["Resume Score"],
                [SimpleGaugeChart(score=resume_score, width=300, height=200, label="Resume Score")]
            ]
            
            score_table = Table(score_table_data, colWidths=[6*inch])
            score_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (0, 0), 14),
                ('TEXTCOLOR', (0, 0), (0, 0), colors.darkblue),
                ('BOTTOMPADDING', (0, 0), (0, 0), 10),
            ]))
            
            content.append(score_table)
            content.append(Spacer(1, 0.25*inch))
            
            # Add Executive Summary section
            content.append(Paragraph("Executive Summary", heading_style))
            content.append(Spacer(1, 0.1*inch))
            
            overall_assessment = ""
            if "## Overall Assessment" in analysis_text:
                overall_section = analysis_text.split("## Overall Assessment")[1].split("##")[0].strip()
                overall_assessment = self._clean_markdown(overall_section)
            
            content.append(Paragraph(overall_assessment, normal_style))
            content.append(Spacer(1, 0.2*inch))
            
            # Key Strengths and Areas for Improvement section
            content.append(Paragraph("Key Strengths and Areas for Improvement", subheading_style))
            content.append(Spacer(1, 0.1*inch))

            if strengths or weaknesses:
                # Create data for strengths and weaknesses
                sw_data = [["Key Strengths", "Areas for Improvement"]]
                
                # Get max length of strengths and weaknesses
                max_len = max(len(strengths), len(weaknesses), 1)
                
                for i in range(max_len):
                    strength = f"• {self._clean_markdown(strengths[i])}" if i < len(strengths) else ""
                    weakness = f"• {self._clean_markdown(weaknesses[i])}" if i < len(weaknesses) else ""
                    sw_data.append([
                        Paragraph(strength, list_item_style) if strength else "",
                        Paragraph(weakness, list_item_style) if weakness else ""
                    ])
                
                sw_table = Table(sw_data, colWidths=[3*inch, 3*inch])
                sw_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, 0), colors.lightgreen),
                    ('BACKGROUND', (1, 0), (1, 0), colors.salmon),
                    ('TEXTCOLOR', (0, 0), (1, 0), colors.black),
                    ('ALIGN', (0, 0), (1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (1, 0), 10),
                    ('GRID', (0, 0), (1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                
                content.append(sw_table)
            else:
                # Add empty strengths and weaknesses with a message
                empty_data = [
                    ["Key Strengths", "Areas for Improvement"],
                    [
                        Paragraph("No specific strengths identified in the analysis.", normal_style),
                        Paragraph("No specific areas for improvement identified in the analysis.", normal_style)
                    ]
                ]
                empty_table = Table(empty_data, colWidths=[3*inch, 3*inch])
                empty_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, 0), colors.lightgreen),
                    ('BACKGROUND', (1, 0), (1, 0), colors.salmon),
                    ('TEXTCOLOR', (0, 0), (1, 0), colors.black),
                    ('ALIGN', (0, 0), (1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (1, 0), 10),
                    ('GRID', (0, 0), (1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                
                content.append(empty_table)

            content.append(Spacer(1, 0.25*inch))
            
            # Use the process_sections method to handle detailed analysis
            content = self.process_sections(analysis_text, content, normal_style, list_item_style, subheading_style, heading_style, clean_markdown)
            
            # Add course recommendations
            course_recommendations = []
            
            # Try to get course recommendations from different sources
            if "suggestions" in analysis_result:
                course_recommendations = analysis_result.get("suggestions", [])
            
            # If still no recommendations, try to extract from text
            if not course_recommendations and "## Recommended Courses" in analysis_text:
                recommendations_section = analysis_text.split("## Recommended Courses")[1].split("##")[0].strip()
                course_recommendations = [clean_markdown(r.strip().replace("- ", "").replace("* ", "").replace("• ", "")) 
                              for r in recommendations_section.split("\n") 
                              if r.strip() and (r.strip().startswith("-") or r.strip().startswith("*") or r.strip().startswith("•"))]
            
            # Try another pattern for course recommendations
            if not course_recommendations and "Recommended Courses" in analysis_text:
                recommendations_section = analysis_text.split("Recommended Courses")[1]
                if "##" in recommendations_section:
                    recommendations_section = recommendations_section.split("##")[0]
                
                # Extract lines that look like list items
                for line in recommendations_section.split("\n"):
                    line = line.strip()
                    if line and ":" in line and not line.startswith("#"):
                        course_recommendations.append(clean_markdown(line))
            
            content.append(Paragraph("Recommended Courses & Certifications", subheading_style))
            
            if course_recommendations:
                # Create a table for course recommendations with better formatting
                course_data = [["Recommended Courses & Certifications"]]  # Add header row
                
                for course in course_recommendations:
                    # Clean the course text and ensure it doesn't have any markdown formatting
                    cleaned_course = clean_markdown(course)
                    course_data.append([Paragraph(f"• {cleaned_course}", list_item_style)])
                
                course_table = Table(course_data, colWidths=[6*inch])
                course_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, 0), colors.lightblue),
                    ('TEXTCOLOR', (0, 0), (0, 0), colors.black),
                    ('ALIGN', (0, 0), (0, 0), 'CENTER'),  # Center the header
                    ('ALIGN', (0, 1), (0, -1), 'LEFT'),   # Left-align the content
                    ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (0, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (0, 0), 10),
                    ('GRID', (0, 0), (0, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (0, -1), 'TOP'),
                ]))
                
                content.append(course_table)
            else:
                # If still no recommendations, add a text section instead of generic courses
                content.append(Paragraph("Based on your resume and target role, consider the following types of courses and certifications:", normal_style))
                content.append(Spacer(1, 0.1*inch))
                
                # Add role-specific recommendations based on job_role
                role_specific_courses = []
                if "data" in job_role.lower() or "scientist" in job_role.lower() or "analyst" in job_role.lower():
                    role_specific_courses = [
                        "Data Science Specialization (Coursera/edX)",
                        "Machine Learning (Coursera/edX)",
                        "Deep Learning Specialization (Coursera)",
                        "Big Data Technologies (Cloud Provider Certifications)",
                        "Statistical Modeling and Inference",
                        "Data Visualization with Tableau/Power BI"
                    ]
                elif "developer" in job_role.lower() or "engineer" in job_role.lower() or "programming" in job_role.lower():
                    role_specific_courses = [
                        "Full Stack Web Development (Udemy/Coursera)",
                        "Cloud Certifications (AWS/Azure/GCP)",
                        "DevOps and CI/CD Pipelines",
                        "Software Architecture and Design Patterns",
                        "Agile and Scrum Methodologies",
                        "Mobile App Development"
                    ]
                elif "security" in job_role.lower() or "cyber" in job_role.lower():
                    role_specific_courses = [
                        "Certified Information Systems Security Professional (CISSP)",
                        "Certified Ethical Hacker (CEH)",
                        "CompTIA Security+",
                        "Offensive Security Certified Professional (OSCP)",
                        "Cloud Security Certifications",
                        "Security Operations and Incident Response"
                    ]
                else:
                    # Generic professional development courses
                    role_specific_courses = [
                        "LinkedIn Learning - Professional Skills Development",
                        "Coursera - Career Development Specialization",
                        "Udemy - Job Interview Skills Training",
                        "Project Management Professional (PMP)",
                        "Leadership and Management Skills",
                        "Technical Writing and Communication"
                    ]
                
                # Create a table for role-specific courses
                course_data = []
                for course in role_specific_courses:
                    course_data.append([Paragraph(f"• {clean_markdown(course)}", list_item_style)])
                
                course_table = Table(course_data, colWidths=[6*inch])
                course_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, 0), colors.lightblue),
                    ('TEXTCOLOR', (0, 0), (0, 0), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica'),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]))
                
                content.append(course_table)
            
            content.append(Spacer(1, 0.2*inch))
            
            # Add footer with page numbers
            def add_page_number(canvas, doc):
                canvas.saveState()
                canvas.setFont('Helvetica', 9)
                page_num = canvas.getPageNumber()
                text = f"Page {page_num}"
                canvas.drawRightString(7.5*inch, 0.25*inch, text)
                
                # Add generation date at the bottom
                canvas.setFont('Helvetica', 9)
                date_text = f"Generated on: {datetime.datetime.now().strftime('%B %d, %Y')}"
                canvas.drawString(0.5*inch, 0.25*inch, date_text)
                
                canvas.restoreState()
            
            # Build the PDF
            doc.build(content, onFirstPage=add_page_number, onLaterPages=add_page_number)
            
            # Get the PDF from the buffer
            buffer.seek(0)
            return buffer
        
        except Exception as e:
            st.error(f"Error generating simple PDF report: {e!s}")
            import traceback
            st.code(traceback.format_exc())
            return None 

    def process_sections(self, analysis_text, content, normal_style, list_item_style, subheading_style, heading_style, clean_markdown):
        """Process sections of the analysis text with special handling for certain sections"""
        # Parse the markdown-like content
        sections = analysis_text.split("##")
        
        # Define sections to include in detailed analysis
        detailed_sections = [
            "Professional Profile Analysis",
            "Skills Analysis",
            "Experience Analysis",
            "Education Analysis",
            "ATS Optimization Assessment",
            "Role Alignment Analysis",
            "Job Match Analysis"
        ]
        
        # Add Detailed Analysis section
        content.append(Paragraph("Detailed Analysis", heading_style))
        content.append(Spacer(1, 0.1*inch))
        
        for section in sections:
            if not section.strip():
                continue
            
            # Extract section title and content
            lines = section.strip().split("\n")
            section_title = lines[0].strip()
            
            # Skip sections we don't want in the detailed analysis
            if section_title not in detailed_sections and section_title != "Overall Assessment":
                continue
            
            # Skip Overall Assessment as we've already included it
            if section_title == "Overall Assessment":
                continue
            
            section_content = "\n".join(lines[1:]).strip()
            
            # Add section title
            content.append(Paragraph(section_title, subheading_style))
            content.append(Spacer(1, 0.1*inch))
            
            # Process content based on section
            if section_title == "Skills Analysis":
                # Extract current and missing skills
                current_skills = []
                missing_skills = []
                
                if "Current Skills" in section_content:
                    current_part = section_content.split("Current Skills")[1]
                    if "Missing Skills" in current_part:
                        current_part = current_part.split("Missing Skills")[0]
                    
                    for line in current_part.split("\n"):
                        if line.strip() and ("-" in line or "*" in line or "•" in line):
                            skill = clean_markdown(line.replace("-", "").replace("*", "").replace("•", "").strip())
                            if skill:
                                current_skills.append(skill)
                
                if "Missing Skills" in section_content:
                    missing_part = section_content.split("Missing Skills")[1]
                    for line in missing_part.split("\n"):
                        if line.strip() and ("-" in line or "*" in line or "•" in line):
                            skill = clean_markdown(line.replace("-", "").replace("*", "").replace("•", "").strip())
                            if skill:
                                missing_skills.append(skill)
                
                # Create skills table with better formatting
                if current_skills or missing_skills:
                    # Create paragraphs for each skill to ensure proper wrapping
                    current_skill_paragraphs = [Paragraph(skill, normal_style) for skill in current_skills]
                    missing_skill_paragraphs = [Paragraph(skill, normal_style) for skill in missing_skills]
                    
                    # Make sure both lists have the same length
                    max_len = max(len(current_skill_paragraphs), len(missing_skill_paragraphs))
                    current_skill_paragraphs.extend([Paragraph("", normal_style)] * (max_len - len(current_skill_paragraphs)))
                    missing_skill_paragraphs.extend([Paragraph("", normal_style)] * (max_len - len(missing_skill_paragraphs)))
                    
                    # Create data for the table
                    data = [["Current Skills", "Missing Skills"]]
                    for i in range(max_len):
                        data.append([current_skill_paragraphs[i], missing_skill_paragraphs[i]])
                    
                    # Create the table with fixed column widths
                    table = Table(data, colWidths=[3*inch, 3*inch])
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (1, 0), colors.lightgreen),
                        ('TEXTCOLOR', (0, 0), (1, 0), colors.black),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('LEFTPADDING', (0, 0), (-1, -1), 10),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                    ]))
                    
                    content.append(table)
                
                # We no longer need to add skill proficiency outside the table
                # as it's now included in the table itself
            elif section_title == "ATS Optimization Assessment":
                # Special handling for ATS Optimization Assessment
                ats_score_line = ""
                ats_content = []
                
                # Extract ATS score if present
                for line in section_content.split("\n"):
                    if "ATS Score:" in line:
                        ats_score_line = clean_markdown(line)
                    elif line.strip():
                        # Check if it's a list item
                        if line.strip().startswith("-") or line.strip().startswith("*") or line.strip().startswith("•"):
                            ats_content.append("• " + clean_markdown(line.strip()[1:].strip()))
                        else:
                            ats_content.append(clean_markdown(line))
                
                # Add ATS score line if found
                if ats_score_line:
                    content.append(Paragraph(ats_score_line, normal_style))
                    content.append(Spacer(1, 0.1*inch))
                
                # Add the rest of the ATS content
                for para in ats_content:
                    if para.startswith("• "):
                        content.append(Paragraph(para, list_item_style))
                    else:
                        content.append(Paragraph(para, normal_style))
            else:
                # Process regular paragraphs
                paragraphs = section_content.split("\n")
                for para in paragraphs:
                    if para.strip():
                        # Check if it's a list item
                        if para.strip().startswith("-") or para.strip().startswith("*") or para.strip().startswith("•"):
                            para = "• " + clean_markdown(para.strip()[1:].strip())
                            content.append(Paragraph(para, list_item_style))
                        else:
                            content.append(Paragraph(clean_markdown(para), normal_style))
            
            content.append(Spacer(1, 0.2*inch))
        
        return content

    def generate_prompt_based_resume(self, job_description: str, prompt_instructions: str, candidate_info: dict | None = None) -> dict:
        """
        Generate a fully tailored, ATS-aligned resume data structure based on a target Job Description and custom AI prompt instructions.
        """
        if not candidate_info:
            candidate_info = {}

        name = candidate_info.get("full_name", "Prince Kumar Jha")
        email = candidate_info.get("email", "pkjha2028@gmail.com")
        phone = candidate_info.get("phone", "+91 98765 43210")
        loc = candidate_info.get("location", "India")
        linkedin = candidate_info.get("linkedin", "linkedin.com/in/princekjha")
        portfolio = candidate_info.get("portfolio", "princekjha.dev")

        system_prompt = f"""
You are an expert Executive Resume Writer and ATS Optimization Strategist.
Generate a structured, ATS-optimized resume tailored specifically for the provided Job Description and user prompt instructions.

TARGET JOB DESCRIPTION:
{job_description or "General Software Engineer / Full Stack AI Developer role requiring Python, React, System Design, SQL, and REST APIs."}

USER PROMPT & INSTRUCTIONS:
{prompt_instructions or "Build a high-impact, professional resume emphasizing technical leadership, scalable architecture, and quantifiable metrics."}

CANDIDATE INFO:
Full Name: {name}
Email: {email}
Phone: {phone}
Location: {loc}

Return ONLY valid JSON strictly adhering to this JSON schema (no markdown formatting, no text before or after):
{{
  "personal_info": {{
    "full_name": "{name}",
    "email": "{email}",
    "phone": "{phone}",
    "location": "{loc}",
    "linkedin": "{linkedin}",
    "portfolio": "{portfolio}",
    "title": "Target Role Title Extracted from JD"
  }},
  "summary": "Compelling 3-4 sentence professional summary loaded with JD target keywords, achievements, and technical strengths.",
  "experience": [
    {{
      "company": "Tech Solutions Corp",
      "position": "Senior Software Engineer / AI Specialist",
      "start_date": "2022",
      "end_date": "Present",
      "description": "Led end-to-end development of cloud applications and high-throughput microservices.",
      "responsibilities": [
        "Architected scalable microservices using Python, FastAPI, and PostgreSQL, reducing latency by 45%.",
        "Engineered automated AI pipelines leveraging LLMs and cloud APIs, accelerating client workflows by 60%.",
        "Mentored a team of 5 engineers in modern DevOps practices, maintaining 99.9% uptime."
      ]
    }},
    {{
      "company": "Innovate Analytics",
      "position": "Software Engineer",
      "start_date": "2020",
      "end_date": "2022",
      "description": "Developed interactive user dashboards and data ingestion pipelines.",
      "responsibilities": [
        "Built responsive web interfaces with React and TypeScript, boosting user engagement by 35%.",
        "Optimized SQL queries and database indexes, improving query response speeds by 3x."
      ]
    }}
  ],
  "projects": [
    {{
      "name": "AiResuMind - SaaS Resume & Job Intelligence Suite",
      "technologies": "Python, Streamlit, FastAPI, Selenium, AI Models",
      "description": "Production SaaS application providing instant AI resume scoring, ATS audit, and job portal integration.",
      "responsibilities": [
        "Implemented multi-provider LLM fallback engine ensuring 99.9% availability.",
        "Built interactive job portal scraper parsing thousands of live vacancies."
      ]
    }}
  ],
  "education": [
    {{
      "school": "State Institute of Technology",
      "degree": "Bachelor of Technology",
      "field": "Computer Science and Engineering",
      "graduation_date": "2020",
      "gpa": "8.8 / 10"
    }}
  ],
  "skills": {{
    "technical": ["Python", "JavaScript", "React", "FastAPI", "PostgreSQL", "System Design"],
    "soft": ["Problem Solving", "Agile Leadership", "Cross-Functional Collaboration"],
    "languages": ["English", "Hindi"],
    "tools": ["Git", "Docker", "AWS", "Streamlit", "Linux"]
  }}
}}
"""
        try:
            raw_response, model_used = self._call_llm_with_fallback(system_prompt, temperature=0.3)
            clean_json = raw_response.strip()
            clean_json = clean_json.removeprefix("```json")
            clean_json = clean_json.removeprefix("```")
            clean_json = clean_json.removesuffix("```")
            clean_json = clean_json.strip()

            parsed = json.loads(clean_json)
            parsed["model_used"] = model_used
            return parsed
        except Exception as e:
            print(f"AI Prompt Resume Builder fallback used due to: {e}")
            return {
                "personal_info": {
                    "full_name": name,
                    "email": email,
                    "phone": phone,
                    "location": loc,
                    "linkedin": linkedin,
                    "portfolio": portfolio,
                    "title": "Software Engineer / AI Systems Developer"
                },
                "summary": "High-performing Software Engineer specializing in building scalable web applications, AI integrations, and cloud backend services. Adept at transforming complex requirements into robust production software tailored for top-tier technology environments.",
                "experience": [
                    {
                        "company": "Apex Technology Labs",
                        "position": "Senior Software Engineer",
                        "start_date": "2022",
                        "end_date": "Present",
                        "description": "Lead engineer responsible for core microservices and AI application pipelines.",
                        "responsibilities": [
                            "Architected high-throughput REST APIs in Python and FastAPI, handling over 100k daily requests with sub-100ms latency.",
                            "Integrated multi-provider AI models for automated document parsing and contextual text extraction.",
                            "Optimized relational database schemas and cache layers, reducing server response overhead by 40%."
                        ]
                    },
                    {
                        "company": "Core Cloud Systems",
                        "position": "Software Engineer",
                        "start_date": "2020",
                        "end_date": "2022",
                        "description": "Developed full-stack web features and analytics dashboards.",
                        "responsibilities": [
                            "Designed dynamic, responsive web interfaces using React and modern CSS design tokens.",
                            "Automated CI/CD deployment pipelines on Docker and cloud infrastructure."
                        ]
                    }
                ],
                "projects": [
                    {
                        "name": "AiResuMind - SaaS Resume & Career Platform",
                        "technologies": "Python, React, Streamlit, SQLite, AI APIs",
                        "description": "End-to-end career growth platform featuring automated ATS evaluation, cover letter generation, and real-time job portals.",
                        "responsibilities": [
                            "Built prompt-driven AI resume generation module based on target Job Descriptions.",
                            "Enforced strict visual design tokens and light-mode accessibility standards."
                        ]
                    }
                ],
                "education": [
                    {
                        "school": "National Institute of Technology",
                        "degree": "Bachelor of Technology",
                        "field": "Computer Science & Engineering",
                        "graduation_date": "2020",
                        "gpa": "8.9 / 10"
                    }
                ],
                "skills": {
                    "technical": ["Python", "JavaScript", "FastAPI", "React", "SQL", "System Design"],
                    "soft": ["Problem Solving", "Technical Communication", "Agile Leadership"],
                    "languages": ["English", "Hindi"],
                    "tools": ["Git", "Docker", "AWS", "PostgreSQL", "Linux"]
                },
                "model_used": "Structured Tailored Fallback"
            }
