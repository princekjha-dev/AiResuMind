import builtins

if not hasattr(builtins, "xrange"):
    builtins.xrange = range

from io import BytesIO

import docx
import pypdf


class ResumeParser:
    def __init__(self):
        pass
        
    def extract_text_from_pdf(self, pdf_file):
        try:
            if hasattr(pdf_file, 'seek'):
                pdf_file.seek(0)
            if hasattr(pdf_file, 'getbuffer'):
                file_content = pdf_file.getbuffer()
            elif hasattr(pdf_file, 'read'):
                file_content = pdf_file.read()
                if hasattr(pdf_file, 'seek'):
                    pdf_file.seek(0)
            else:
                file_content = pdf_file
                
            pdf_reader = pypdf.PdfReader(BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text.strip()
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return ""
            
    def extract_text_from_docx(self, docx_file):
        try:
            if hasattr(docx_file, 'seek'):
                docx_file.seek(0)
            if hasattr(docx_file, 'getbuffer'):
                file_content = docx_file.getbuffer()
            elif hasattr(docx_file, 'read'):
                file_content = docx_file.read()
                if hasattr(docx_file, 'seek'):
                    docx_file.seek(0)
            elif isinstance(docx_file, (bytes, bytearray)):
                file_content = docx_file
            else:
                file_content = bytes(docx_file)

            doc = docx.Document(BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text += paragraph.text.strip() + "\n"
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            return text.strip()
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return ""
            
    def extract_text(self, file):
        if hasattr(file, 'seek'):
            file.seek(0)
        
        fname = getattr(file, 'name', '').lower()
        if fname.endswith('.pdf'):
            return self.extract_text_from_pdf(file)
        elif fname.endswith('.docx') or fname.endswith('.doc'):
            return self.extract_text_from_docx(file)
        else:
            try:
                if hasattr(file, 'getvalue'):
                    return file.getvalue().decode('utf-8', errors='ignore')
                elif hasattr(file, 'read'):
                    content = file.read()
                    if isinstance(content, bytes):
                        return content.decode('utf-8', errors='ignore')
                    return str(content)
            except Exception:
                pass
            return ""

    def parse(self, file):
        text = self.extract_text(file)
        
        skills = []
        experience = []
        education = []
        
        skill_keywords = ['python', 'java', 'javascript', 'html', 'css', 'sql', 'react', 'angular', 'vue', 
                         'node', 'express', 'django', 'flask', 'spring', 'docker', 'kubernetes', 'aws', 
                         'azure', 'git', 'jenkins', 'jira']
                         
        text_lower = text.lower()
        for skill in skill_keywords:
            if skill in text_lower:
                skills.append(skill)
                
        return {
            "skills": skills,
            "experience": experience,
            "education": education,
            "raw_text": text
        }

# Alias for backward compatibility
FileParser = ResumeParser()